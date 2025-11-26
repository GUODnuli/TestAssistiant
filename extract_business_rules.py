#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从Word文档中提取业务规则的脚本
"""
import docx
import sys
import os

def extract_text_from_docx(file_path):
    """
    从DOCX文件中提取文本内容
    
    Args:
        file_path (str): DOCX文件路径
        
    Returns:
        str: 提取的文本内容
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 只添加非空段落
                full_text.append(paragraph.text.strip())
                
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append('\t'.join(row_text))
                    
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"读取DOCX文件时出错: {str(e)}")

def main():
    if len(sys.argv) != 2:
        print("使用方法: python extract_business_rules.py <word_file_path>")
        sys.exit(1)
        
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"错误: 文件 {file_path} 不存在")
        sys.exit(1)
        
    try:
        content = extract_text_from_docx(file_path)
        print(content)
    except Exception as e:
        print(f"错误: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()